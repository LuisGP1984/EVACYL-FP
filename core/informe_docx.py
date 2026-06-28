"""
Generación del informe individual de un alumno en Word (.docx), a partir
de las estructuras recopiladas en core/informe_alumno.py. Incluye la
nota por RA y el detalle por criterio.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.informe_alumno import InformeAlumno, InformeAlumnoCompleto, InformeAlumnoFinal

COLOR_CABECERA = RGBColor(0xE2, 0x2B, 0x10)
COLOR_RESULTADO = RGBColor(0x7A, 0x1B, 0x08)


def _formatear_numero(valor: float | None) -> str:
    if valor is None:
        return "—"
    return f"{valor:.2f}".replace(".", ",")


def _sombrear_celda(celda, color_hex: str):
    propiedades_celda = celda._tc.get_or_add_tcPr()
    sombreado = propiedades_celda.makeelement(qn("w:shd"), {qn("w:fill"): color_hex})
    propiedades_celda.append(sombreado)


def _anadir_tabla(documento: Document, encabezados: list[str], filas: list[list[str]]):
    tabla = documento.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"
    fila_cabecera = tabla.rows[0]
    for celda, texto in zip(fila_cabecera.cells, encabezados):
        celda.text = texto
        _sombrear_celda(celda, "E22B10")
        for parrafo in celda.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in parrafo.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for indice_fila, datos_fila in enumerate(filas):
        fila_tabla = tabla.add_row()
        if indice_fila % 2 == 1:
            for celda in fila_tabla.cells:
                _sombrear_celda(celda, "FDF3E7")
        for celda, valor in zip(fila_tabla.cells, datos_fila):
            celda.text = str(valor)
            for parrafo in celda.paragraphs:
                parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tabla


def _anadir_cabecera_comun(
    documento: Document, nombre_modulo: str, etiqueta_evaluacion: str, apellidos: str, nombre: str
):
    titulo = documento.add_heading("Informe de calificaciones", level=0)
    for run in titulo.runs:
        run.font.color.rgb = COLOR_RESULTADO

    parrafo_alumno = documento.add_paragraph()
    parrafo_alumno.add_run("Alumno/a: ").bold = True
    parrafo_alumno.add_run(f"{apellidos}, {nombre}")

    parrafo_modulo = documento.add_paragraph()
    parrafo_modulo.add_run("Módulo: ").bold = True
    parrafo_modulo.add_run(nombre_modulo)

    parrafo_evaluacion = documento.add_paragraph()
    parrafo_evaluacion.add_run("Evaluación: ").bold = True
    parrafo_evaluacion.add_run(etiqueta_evaluacion)

    documento.add_paragraph()


def _anadir_nota_final(documento: Document, etiqueta: str, valor: float | None, calificacion: str):
    parrafo = documento.add_paragraph()
    run = parrafo.add_run(f"{etiqueta}: {_formatear_numero(valor)} ({calificacion or '—'})")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_RESULTADO
    documento.add_paragraph()


def _cuerpo_seccion_evaluacion_docx(documento: Document, informe: InformeAlumno):
    documento.add_heading("Calificación por Resultado de Aprendizaje", level=1)
    filas_ra = [
        [f"RA{fila.numero_ra}", f"{fila.peso_en_modulo:g}%", _formatear_numero(fila.valor), fila.calificacion or "—"]
        for fila in informe.filas_ra
    ]
    _anadir_tabla(documento, ["RA", "Peso en el módulo", "Nota", "Calificación"], filas_ra)

    documento.add_paragraph()
    documento.add_heading("Detalle por criterio de evaluación", level=1)
    filas = [
        [fila.codigo_criterio, f"{fila.peso_en_ra:g}", _formatear_numero(fila.valor), fila.calificacion or "—"]
        for fila in informe.filas_criterios
    ]
    _anadir_tabla(documento, ["Criterio", "Peso relativo en su RA", "Nota", "Calificación"], filas)

    documento.add_paragraph()
    documento.add_heading("Desglose por instrumento de evaluación", level=1)
    if not informe.bloques_instrumentos:
        documento.add_paragraph("No hay instrumentos de evaluación con criterios asignados.")
    for bloque in informe.bloques_instrumentos:
        parrafo_bloque = documento.add_paragraph()
        parrafo_bloque.add_run(f"{bloque.nombre_instrumento}").bold = True
        parrafo_bloque.add_run(f" — peso {bloque.peso_global:g}% de la evaluación")
        filas_bloque = [
            [
                fila.codigo_criterio,
                f"{fila.peso_en_instrumento:g}%",
                _formatear_numero(fila.valor),
                fila.calificacion or "—",
            ]
            for fila in bloque.criterios
        ]
        _anadir_tabla(documento, ["Criterio", "Peso en este instrumento", "Nota", "Calificación"], filas_bloque)
        documento.add_paragraph()


def generar_informe_evaluacion_docx(informe: InformeAlumno, ruta_destino: str | Path) -> Path:
    ruta_destino = Path(ruta_destino)
    documento = Document()

    _anadir_cabecera_comun(
        documento, informe.nombre_modulo, informe.nombre_evaluacion, informe.apellidos_alumno, informe.nombre_alumno
    )
    _anadir_nota_final(
        documento, "Nota final de la evaluación", informe.nota_final_numerica, informe.calificacion_final
    )

    _cuerpo_seccion_evaluacion_docx(documento, informe)

    documento.save(str(ruta_destino))
    return ruta_destino


def _cuerpo_seccion_final_docx(documento: Document, informe: InformeAlumnoFinal):
    pesos_texto = "  ·  ".join(f"{nombre}: {peso:g}" for nombre, peso in informe.pesos_evaluaciones.items())
    parrafo_pesos = documento.add_paragraph()
    run_pesos = parrafo_pesos.add_run(f"Pesos aplicados entre evaluaciones — {pesos_texto}")
    run_pesos.italic = True
    documento.add_paragraph()

    documento.add_heading("Calificación final por Resultado de Aprendizaje", level=1)
    filas_ra = [
        [f"RA{fila.numero_ra}", f"{fila.peso_en_modulo:g}%", _formatear_numero(fila.valor), fila.calificacion or "—"]
        for fila in informe.filas_ra
    ]
    _anadir_tabla(documento, ["RA", "Peso en el módulo", "Nota FINAL", "Calificación"], filas_ra)

    documento.add_paragraph()
    documento.add_heading("Calificación de cada criterio, por evaluación", level=1)
    encabezados_criterios = ["Criterio"] + informe.nombres_evaluaciones + ["Nota FINAL", "Calif."]
    filas = []
    for fila in informe.filas_criterios:
        valores_formateados = [_formatear_numero(v) for v in fila.valores_por_evaluacion]
        filas.append(
            [fila.codigo_criterio] + valores_formateados + [_formatear_numero(fila.valor), fila.calificacion or "—"]
        )
    _anadir_tabla(documento, encabezados_criterios, filas)

    documento.add_paragraph()
    parrafo_nota = documento.add_paragraph(
        "Las celdas en blanco (—) indican que ese criterio no tuvo ninguna nota calculada en esa "
        "evaluación. La nota FINAL combina las evaluaciones que sí tienen nota, según los pesos "
        "indicados arriba."
    )
    parrafo_nota.runs[0].italic = True


def generar_informe_final_docx(informe: InformeAlumnoFinal, ruta_destino: str | Path) -> Path:
    ruta_destino = Path(ruta_destino)
    documento = Document()

    _anadir_cabecera_comun(
        documento, informe.nombre_modulo, "Evaluación final de curso",
        informe.apellidos_alumno, informe.nombre_alumno,
    )
    _anadir_nota_final(documento, "Nota final de curso", informe.nota_final_numerica, informe.calificacion_final)

    _cuerpo_seccion_final_docx(documento, informe)

    documento.save(str(ruta_destino))
    return ruta_destino


def generar_informe_completo_docx(informe: "InformeAlumnoCompleto", ruta_destino: str | Path) -> Path:
    """Genera un único documento Word con una sección por cada
    evaluación parcial evaluable, más una sección FINAL al final —
    todo en el mismo documento, separado por saltos de página.
    """
    ruta_destino = Path(ruta_destino)
    documento = Document()

    _anadir_cabecera_comun(
        documento, informe.nombre_modulo, "Informe completo del curso",
        informe.apellidos_alumno, informe.nombre_alumno,
    )

    for indice, seccion in enumerate(informe.secciones_evaluaciones):
        if indice > 0:
            documento.add_page_break()
        documento.add_heading(f"Evaluación: {seccion.nombre_evaluacion}", level=0)
        _anadir_nota_final(
            documento, "Nota final de la evaluación", seccion.nota_final_numerica, seccion.calificacion_final
        )
        _cuerpo_seccion_evaluacion_docx(documento, seccion)

    if informe.seccion_final is not None:
        documento.add_page_break()
        documento.add_heading("Evaluación FINAL", level=0)
        _anadir_nota_final(
            documento, "Nota final de curso",
            informe.seccion_final.nota_final_numerica, informe.seccion_final.calificacion_final,
        )
        _cuerpo_seccion_final_docx(documento, informe.seccion_final)

    documento.save(str(ruta_destino))
    return ruta_destino
